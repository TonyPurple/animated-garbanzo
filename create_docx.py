from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('Review Service API Design', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add Assumptions section
doc.add_heading('Assumptions', level=2)
assumptions = [
    "The API assumes authentication and authorization are handled at a different layer (e.g., API gateway or middleware).",
    "All requests reaching the API are pre-validated for authentication and authorization.",
    "Users can submit only one review per product, identified by their `userId`, but can update their existing review.",
    "Reviews will be moderated externally, and no content filtering is applied at the API level.",
    "Reviews belong to products, so they are nested under `/products/{id}/reviews` for creation and listing.",
    "Individual review actions (fetch, update, delete) use `/reviews/{id}` as reviews have unique IDs.",
    "The database structure follows a NoSQL schema to store reviews in a flexible format."
]

for assumption in assumptions:
    doc.add_paragraph(assumption, style='List Bullet')

# Add Database Choice section
doc.add_heading('Database Choice: MongoDB', level=2)
doc.add_heading('Reasoning:', level=3)
reasoning = [
    "Flexible Schema: Reviews can have optional fields such as ratings, text comments, and metadata.",
    "Scalability: MongoDB is optimized for large-scale applications with high read/write operations.",
    "Document-based Storage: Each review can be stored as a document, making it easy to retrieve by product ID or user ID."
]

for reason in reasoning:
    doc.add_paragraph(reason, style='List Bullet')

# Add Database Schema Section
doc.add_heading('Database Schema', level=2)
schema_code = '''
// Reviews Collection

{
  _id: ObjectId, // MongoDB's default unique ID
  productId: String, // ID of the product being reviewed
  userId: String, // unique user id
  rating: Number, // Optional: Rating (e.g., 1-5)
  comment: String, // Optional: Text comment
  reviewerName: String,
  reviewerEmail: String,
  metadata: Object, // Optional: Additional metadata
  createdAt: Date, // Timestamp of review creation
  updatedAt: Date // Optional: Timestamp of review update
}
'''
schema_paragraph = doc.add_paragraph()
schema_paragraph.add_run(schema_code).font.name = 'Courier New'

# Add API Endpoints section
doc.add_heading('API Endpoints', level=2)

# Endpoint 1: Create a Review
doc.add_heading('1. Create a Review', level=3)
doc.add_paragraph('Method: POST')
doc.add_paragraph('Endpoint: /products/{id}/reviews')
doc.add_heading('Request Payload:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "userId": "user123",
    "rating": 5,
    "comment": "Great product!",
    "reviewerName": "John Doe",
    "reviewerEmail": "john.doe@example.com",
    "metadata": {"key": "value"}
}''').font.name = 'Courier New'

doc.add_heading('Response Payload:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "id": "unique_review_id",
    "productId": "{id}",
    "userId": "user123",
    "rating": 5,
    "comment": "Great product!",
    "reviewerName": "John Doe",
    "reviewerEmail": "john.doe@example.com",
    "metadata": {"key": "value"},
    "createdAt": "2024-10-27T10:00:00Z"
}''').font.name = 'Courier New'

doc.add_heading('Error Responses:', level=4)
doc.add_paragraph('400 Bad Request:', style='List Bullet')
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "error": "Bad Request",
    "message": "The request payload is invalid. Must contain at least one of rating or comment."
}''').font.name = 'Courier New'

doc.add_paragraph('409 Conflict:', style='List Bullet')
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "error": "Conflict",
    "message": "A review already exists for this product by the same user."
}''').font.name = 'Courier New'

# Endpoint 2: Get Reviews for a Product
doc.add_heading('2. Get Reviews for a Product', level=3)
doc.add_paragraph('Method: GET')
doc.add_paragraph('Endpoint: /products/{id}/reviews')
doc.add_heading('Query Parameters:', level=4)
doc.add_paragraph('- page (optional, default: 1) → Pagination support', style='List Bullet')
doc.add_paragraph('- pageSize (optional, default: 10) → Limits the number of reviews returned', style='List Bullet')
doc.add_heading('Example Request:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('GET /products/{id}/reviews?page=2&pageSize=5').font.name = 'Courier New'

doc.add_heading('Response Payload:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
  "reviews": [
    {
      "id": "unique_review_id_6",
      "productId": "{id}",
      "userId": "user123",
      "rating": 5,
      "comment": "Great quality!",
      "reviewerName": "Alice Brown",
      "createdAt": "2024-10-25T14:00:00Z",
      "metadata": {}
    },
    {
      "id": "unique_review_id_7",
      "productId": "{id}",
      "userId": "user456",
      "rating": 3,
      "comment": "Average product.",
      "reviewerName": "Bob White",
      "createdAt": "2024-10-24T16:30:00Z",
      "metadata": {}
    },
    {
      "id": "unique_review_id_8",
      "productId": "{id}",
      "userId": "user789",
      "rating": 4,
      "comment": "Would buy again.",
      "reviewerName": "Charlie Green",
      "createdAt": "2024-10-23T18:15:00Z",
      "metadata": {}
    },
    {
      "id": "unique_review_id_9",
      "productId": "{id}",
      "userId": "user101",
      "rating": 2,
      "comment": "Not as expected.",
      "reviewerName": "David Black",
      "createdAt": "2024-10-22T19:45:00Z",
      "metadata": {}
    },
    {
      "id": "unique_review_id_10",
      "productId": "{id}",
      "userId": "user112",
      "rating": 5,
      "comment": "Loved it!",
      "reviewerName": "Emily White",
      "createdAt": "2024-10-21T21:30:00Z",
      "metadata": {}
    }
  ],
  "total": 15,
  "currentPage": 2,
  "totalPages": 3
}''').font.name = 'Courier New'

doc.add_heading('Error Responses:', level=4)
doc.add_paragraph('404 Not Found:', style='List Bullet')
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "error": "Product not found",
    "message": "The product with ID '123' does not exist."
}''').font.name = 'Courier New'

# Endpoint 3: Get a Single Review by ID
doc.add_heading('3. Get a Single Review by ID', level=3)
doc.add_paragraph('Method: GET')
doc.add_paragraph('Endpoint: /reviews/{id}')
doc.add_heading('Example Request:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('GET /reviews/{id}').font.name = 'Courier New'

doc.add_heading('Response Payload:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "id": "unique_review_id",
    "productId": "{id}",
    "userId": "user123",
    "rating": 5,
    "comment": "Great product!",
    "reviewerName": "John Doe",
    "reviewerEmail": "john.doe@example.com",
    "createdAt": "2024-10-27T10:00:00Z",
    "metadata": {}
}''').font.name = 'Courier New'

doc.add_heading('Error Responses:', level=4)
doc.add_paragraph('404 Not Found:', style='List Bullet')
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "error": "Review not found",
    "message": "The review with ID '123' does not exist."
}''').font.name = 'Courier New'

# Endpoint 4: Update a Review
doc.add_heading('4. Update a Review', level=3)
doc.add_paragraph('Method: PUT')
doc.add_paragraph('Endpoint: /reviews/{id}')
doc.add_heading('Request Payload:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "rating": 4,
    "comment": "Updated comment.",
    "metadata": {"updatedKey":"updatedValue"}
}''').font.name = 'Courier New'

doc.add_heading('Response Payload:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "id": "unique_review_id",
    "productId": "{id}",
    "userId": "user123",
    "rating": 4,
    "comment": "Updated comment.",
    "reviewerName": "John Doe",
    "reviewerEmail": "john.doe@example.com",
    "createdAt": "2024-10-27T10:00:00Z",
    "updatedAt": "2024-10-27T11:00:00Z",
    "metadata": {"updatedKey":"updatedValue"}
}''').font.name = 'Courier New'

doc.add_heading('Error Responses:', level=4)
doc.add_paragraph('404 Not Found:', style='List Bullet')
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "error": "Review not found",
    "message": "The review with ID '123' does not exist."
}''').font.name = 'Courier New'

doc.add_paragraph('400 Bad Request:', style='List Bullet')
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "error": "Bad Request",
    "message": "The request payload is invalid. Must contain at least one of rating or comment."
}''').font.name = 'Courier New'

# Endpoint 5: Delete a Review
doc.add_heading('5. Delete a Review', level=3)
doc.add_paragraph('Method: DELETE')
doc.add_paragraph('Endpoint: /reviews/{id}')
doc.add_heading('Example Request:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('DELETE /reviews/{id}').font.name = 'Courier New'

doc.add_heading('Response Payload:', level=4)
code_paragraph = doc.add_paragraph()
code_paragraph.add_run('''{
    "message": "Review deleted successfully",
    "id": "123",
    "deletedAt": "2024-10-27T12:00:00Z"
}''').font.name = 'Courier New'

# Save the document
doc.save('Review_Service_API_Design.docx')
